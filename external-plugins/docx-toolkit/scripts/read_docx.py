#!/usr/bin/env python3
"""
read_docx.py
Reads an existing .docx and prints its content as Markdown or JSON.

Usage:
    python3 read_docx.py input.docx --format markdown
    python3 read_docx.py input.docx --format json
    python3 read_docx.py input.docx --extract-images images/
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.oxml.ns import qn


def paragraph_to_markdown(paragraph):
    text = paragraph.text
    if not text.strip():
        return ""
    style_obj = paragraph.style
    style = ((style_obj.name if style_obj is not None else None) or "").lower()
    if style.startswith("heading 1"):
        return f"# {text}"
    if style.startswith("heading 2"):
        return f"## {text}"
    if style.startswith("heading 3"):
        return f"### {text}"
    if style.startswith("heading 4"):
        return f"#### {text}"
    if style.startswith("list bullet"):
        return f"- {text}"
    if style.startswith("list number"):
        return f"1. {text}"
    if style.startswith("list paragraph"):
        # Generic list style used by some generators (e.g. docx-js) that don't
        # distinguish bullet vs numbered by style name alone. Render as a
        # bullet by default -- still correct content, just not renumbered.
        return f"- {text}"
    if style == "title":
        return f"# {text}"
    return text


def table_to_markdown(table):
    lines = []
    rows = table.rows
    if not rows:
        return ""
    header = [c.text.strip() for c in rows[0].cells]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def iter_block_items(doc):
    """Yield paragraphs and tables in the order they appear in the document body."""
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def to_markdown(doc):
    out_lines = []
    hf = extract_headers_footers(doc)
    if hf["headers"]:
        out_lines.append(f"> **Header:** {hf['headers'][0]}")
        out_lines.append("")
    for block in iter_block_items(doc):
        if block.__class__.__name__ == "Paragraph":
            md = paragraph_to_markdown(block)
            if md:
                out_lines.append(md)
        else:
            out_lines.append(table_to_markdown(block))
        out_lines.append("")
    if hf["footers"]:
        out_lines.append(f"> **Footer:** {hf['footers'][0]}")
    links = extract_hyperlinks(doc)
    if links:
        out_lines.append("")
        out_lines.append("**Links found:**")
        for link in links:
            out_lines.append(f"- {link['url']}")
    return "\n".join(out_lines).strip() + "\n"


def to_json_struct(doc):
    result = {"headings": [], "paragraphs": [], "tables": []}
    for block in iter_block_items(doc):
        if block.__class__.__name__ == "Paragraph":
            text = block.text
            if not text.strip():
                continue
            style_obj = block.style
            style = (style_obj.name if style_obj is not None else None) or "Normal"
            entry = {"text": text, "style": style}
            result["paragraphs"].append(entry)
            if style.lower().startswith("heading") or style.lower() == "title":
                result["headings"].append(entry)
        else:
            rows = [[c.text.strip() for c in row.cells] for row in block.rows]
            result["tables"].append(rows)
    hf = extract_headers_footers(doc)
    result["headers"] = hf["headers"]
    result["footers"] = hf["footers"]
    result["hyperlinks"] = extract_hyperlinks(doc)
    # Expose core document metadata (title, author, subject, etc.)
    cp = doc.core_properties
    result["metadata"] = {
        "title": cp.title,
        "author": cp.author,
        "subject": cp.subject,
        "keywords": cp.keywords,
        "category": cp.category,
        "comments": cp.comments,
        "language": cp.language,
        "last_modified_by": cp.last_modified_by,
    }
    return result


def extract_images(doc, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    count = 0
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_part = rel.target_part
            ext = image_part.content_type.split("/")[-1]
            count += 1
            fname = os.path.join(out_dir, f"image_{count}.{ext}")
            with open(fname, "wb") as f:
                f.write(image_part.blob)
    return count


def extract_headers_footers(doc):
    result = {"headers": [], "footers": []}
    for section in doc.sections:
        h_text = section.header.paragraphs[0].text if section.header.paragraphs else ""
        f_text = section.footer.paragraphs[0].text if section.footer.paragraphs else ""
        if h_text:
            result["headers"].append(h_text)
        if f_text:
            result["footers"].append(f_text)
    return result


def extract_hyperlinks(doc):
    links = []
    rels = doc.part.rels
    for rel_id, rel in rels.items():
        if "hyperlink" in rel.reltype and rel.is_external:
            links.append({"id": rel_id, "url": rel.target_ref})
    return links


def main():
    parser = argparse.ArgumentParser(description="Read/extract content from a .docx file")
    parser.add_argument("input", help="Path to the .docx file")
    parser.add_argument("--format", choices=["markdown", "json"], default="markdown")
    parser.add_argument("--extract-images", metavar="DIR", default=None,
                         help="Also dump embedded images to this directory")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"error: file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    doc = Document(args.input)

    if args.extract_images:
        n = extract_images(doc, args.extract_images)
        print(f"[read_docx] extracted {n} image(s) to {args.extract_images}", file=sys.stderr)

    if args.format == "json":
        print(json.dumps(to_json_struct(doc), indent=2, ensure_ascii=False))
    else:
        print(to_markdown(doc))


if __name__ == "__main__":
    main()
