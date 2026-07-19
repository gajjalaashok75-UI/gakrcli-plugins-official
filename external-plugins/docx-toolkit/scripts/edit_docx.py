#!/usr/bin/env python3
"""
edit_docx.py
Applies one or more edit operations to an existing .docx and writes a new file.

Usage:
    python3 edit_docx.py input.docx --out output.docx --ops ops.json
    python3 edit_docx.py input.docx --out output.docx --find "old text" --replace "new text"

See SKILL.md section 4 for the ops JSON format.
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Emu, Inches
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def merge_paragraph_runs(paragraph):
    """
    Word frequently splits one visible sentence across several <w:r> runs
    (spell-check boundaries, revision markers, etc). To make find/replace
    reliable, collapse all runs in a paragraph into the first run, carrying
    over the first run's formatting, then drop the rest.
    """
    runs = paragraph.runs
    if len(runs) <= 1:
        return
    full_text = "".join(r.text for r in runs)
    runs[0].text = full_text
    for r in runs[1:]:
        r.text = ""


def op_find_replace(doc, op):
    find = op["find"]
    replace = op["replace"]
    count = 0
    for paragraph in doc.paragraphs:
        if find in paragraph.text:
            merge_paragraph_runs(paragraph)
            for run in paragraph.runs:
                if find in run.text:
                    run.text = run.text.replace(find, replace)
                    count += 1
    # also check inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if find in paragraph.text:
                        merge_paragraph_runs(paragraph)
                        for run in paragraph.runs:
                            if find in run.text:
                                run.text = run.text.replace(find, replace)
                                count += 1
    print(f"[edit_docx] find_replace: {count} occurrence(s) of {find!r} replaced", file=sys.stderr)


def op_append_paragraph(doc, op):
    """
    Appends a paragraph. Supports either plain text (+bold/italic shorthand)
    or a "runs" list for mixed formatting within one paragraph, including
    font color and highlight color.
    """
    p = doc.add_paragraph()
    if "runs" in op:
        for r in op["runs"]:
            run = p.add_run(r.get("text", ""))
            run.bold = bool(r.get("bold"))
            run.italic = bool(r.get("italic"))
            if r.get("color"):
                run.font.color.rgb = RGBColor.from_string(r["color"].replace("#", ""))
            if r.get("highlight"):
                name = r["highlight"].upper()
                if hasattr(WD_COLOR_INDEX, name):
                    run.font.highlight_color = getattr(WD_COLOR_INDEX, name)
    else:
        run = p.add_run(op.get("text", ""))
        run.bold = bool(op.get("bold"))
        run.italic = bool(op.get("italic"))
        if op.get("color"):
            run.font.color.rgb = RGBColor.from_string(op["color"].replace("#", ""))
        if op.get("highlight"):
            name = op["highlight"].upper()
            if hasattr(WD_COLOR_INDEX, name):
                run.font.highlight_color = getattr(WD_COLOR_INDEX, name)


def add_hyperlink_run(paragraph, url, text, color="0563C1", underline=True):
    """
    python-docx has no native API for hyperlink runs, so this builds the
    <w:hyperlink> element directly from OXML primitives: register an external
    relationship for the URL, then construct a run inside a w:hyperlink wrapper
    with the standard blue/underlined hyperlink formatting.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    new_run.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def op_insert_hyperlink(doc, op):
    p = doc.add_paragraph()
    add_hyperlink_run(p, op["url"], op.get("text", op["url"]))


def op_set_header(doc, op):
    text = op.get("text", "")
    section = doc.sections[0]
    section.header.paragraphs[0].text = text


def op_set_footer(doc, op):
    text = op.get("text", "")
    section = doc.sections[0]
    section.footer.paragraphs[0].text = text


def _has_page_break(paragraph):
    """
    Detect an explicit <w:br w:type="page"/> within a paragraph's runs.
    Note: python-docx's built-in `paragraph.contains_page_break` only reflects
    Word's *rendered* page breaks (computed on open in Word), not explicit
    breaks we insert ourselves before the file has ever been opened in Word --
    so that property is unreliable here and this checks the raw run XML instead.
    """
    for run in paragraph.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _get_pages(doc):
    """
    Splits the document body into "pages": lists of top-level block XML
    elements (paragraphs, tables), where a page boundary is any paragraph
    containing an explicit page-break run. The paragraph holding the break
    is included as the last element of the page it closes.
    """
    from docx.text.paragraph import Paragraph

    pages = []
    current = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue  # trailing section properties -- not page content
        current.append(child)
        if child.tag == qn("w:p") and _has_page_break(Paragraph(child, doc)):
            pages.append(current)
            current = []
    if current:
        pages.append(current)
    return pages


def _rebuild_body(doc, pages):
    """Clears the body of block content (keeping sectPr) and re-appends pages in order."""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body.iterchildren()):
        body.remove(child)
    for page in pages:
        for elem in page:
            body.append(elem)
    if sect_pr is not None:
        body.append(sect_pr)


def op_remove_page(doc, op):
    index = op["index"]
    pages = _get_pages(doc)
    if index >= len(pages):
        raise ValueError(f"page index {index} out of range (document has {len(pages)} page(s))")
    del pages[index]
    _rebuild_body(doc, pages)
    print(f"[edit_docx] removed page {index}; {len(pages)} page(s) remain", file=sys.stderr)


def op_move_page(doc, op):
    from_index = op["from"]
    to_index = op["to"]
    pages = _get_pages(doc)
    if from_index >= len(pages) or to_index >= len(pages):
        raise ValueError(f"page index out of range (document has {len(pages)} page(s))")
    page = pages.pop(from_index)
    pages.insert(to_index, page)
    _rebuild_body(doc, pages)
    print(f"[edit_docx] moved page {from_index} -> {to_index}", file=sys.stderr)


def op_resize_image(doc, op):
    """
    Resize the Nth inline image in the document (0-indexed in document order).
    Width/height are in pixels, converted to EMU (1px = 9525 EMU), matching
    the same approximation create_docx.js uses so sizes are consistent
    whether an image was created or later resized.
    """
    index = op.get("index", 0)
    width = op.get("width")
    height = op.get("height")
    shapes = doc.inline_shapes
    if index >= len(shapes):
        raise ValueError(f"image index {index} out of range (found {len(shapes)} images)")
    shape = shapes[index]
    if width:
        shape.width = Emu(int(width * 9525))
    if height:
        shape.height = Emu(int(height * 9525))


def _style_exists(doc, style_name):
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def op_append_heading(doc, op):
    text = op.get("text", "")
    level = op.get("level", 1)
    style_name = f"Heading {level}"
    if _style_exists(doc, style_name):
        doc.add_heading(text, level=level)
    else:
        # The source file (e.g. one produced by a different docx library, such as
        # docx-js) may not define a "Heading N" style under that exact name.
        # python-docx's add_heading() would insert the paragraph and THEN fail
        # when assigning the missing style, leaving a stray duplicate behind --
        # so we check availability up front and fall back to manual formatting
        # that looks like a heading instead.
        sizes = {1: 20, 2: 16, 3: 13, 4: 12}
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 12))


def op_append_table(doc, op):
    rows = op["rows"]
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass  # style not present in this template; leave default
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(cell_text)
            if r_idx == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def op_append_image(doc, op):
    path = op["path"]
    if not os.path.exists(path):
        raise FileNotFoundError(f"image not found: {path}")
    width = op.get("width")
    height = op.get("height")
    from docx.shared import Emu
    kwargs = {}
    if width:
        kwargs["width"] = Emu(width * 9525)  # px -> EMU approximation
    if height:
        kwargs["height"] = Emu(height * 9525)
    doc.add_picture(path, **kwargs)


def op_insert_page_break(doc, op):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _find_paragraph_containing(doc, needle):
    """Locate the first paragraph (body or table cell) whose text contains needle."""
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if needle in paragraph.text:
                        return paragraph
    return None


def op_add_comment(doc, op):
    """
    Anchor a review comment to the run(s) containing the given text.
    Uses python-docx's native comment support (>=1.1), which manages the
    comments part, relationships, and content-type registration internally --
    no manual OOXML wiring required.
    """
    anchor_text = op["anchor"]
    comment_text = op.get("text", "")
    author = op.get("author", "Agent")
    initials = op.get("initials", "A")

    paragraph = _find_paragraph_containing(doc, anchor_text)
    if paragraph is None:
        raise ValueError(f"could not find text to anchor comment: {anchor_text!r}")

    merge_paragraph_runs(paragraph)  # ensure the anchor text lives in a single run
    target_runs = [r for r in paragraph.runs if anchor_text in r.text]
    if not target_runs:
        raise ValueError(f"anchor text not found in any run after merge: {anchor_text!r}")

    doc.add_comment(target_runs, text=comment_text, author=author, initials=initials)


# ---- Metadata ----

def op_set_metadata(doc, op):
    props = doc.core_properties
    field_map = {
        "title": "title", "author": "author", "subject": "subject",
        "keywords": "keywords", "category": "category", "comments": "comments",
        "language": "language", "last_modified_by": "last_modified_by",
    }
    for key, attr in field_map.items():
        if key in op:
            setattr(props, attr, op[key])


# ---- Paragraph formatting ----

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def op_format_paragraph(doc, op):
    """Applies alignment/indent/spacing/keep-with-next to the paragraph containing `anchor` text."""
    anchor = op["anchor"]
    paragraph = _find_paragraph_containing(doc, anchor)
    if paragraph is None:
        raise ValueError(f"could not find paragraph containing: {anchor!r}")
    pf = paragraph.paragraph_format
    if "align" in op:
        pf.alignment = _ALIGN_MAP.get(op["align"])
    indent = op.get("indent", {})
    if "left" in indent:
        pf.left_indent = Inches(indent["left"])
    if "right" in indent:
        pf.right_indent = Inches(indent["right"])
    if "firstLine" in indent:
        pf.first_line_indent = Inches(indent["firstLine"])
    if "hanging" in indent:
        pf.first_line_indent = Inches(-indent["hanging"])
    spacing = op.get("spacing", {})
    if "before" in spacing:
        pf.space_before = Pt(spacing["before"])
    if "after" in spacing:
        pf.space_after = Pt(spacing["after"])
    if "line" in spacing:
        pf.line_spacing = spacing["line"]
    if op.get("keepWithNext"):
        pf.keep_with_next = True
    if op.get("pageBreakBefore"):
        pf.page_break_before = True


# ---- Table structure ----

def _get_table(doc, index):
    tables = doc.tables
    if index >= len(tables):
        raise ValueError(f"table index {index} out of range (found {len(tables)} table(s))")
    return tables[index]


def op_add_table_row(doc, op):
    table = _get_table(doc, op.get("table_index", 0))
    row = table.add_row()
    cells = op.get("cells")
    if cells:
        for i, text in enumerate(cells):
            if i < len(row.cells):
                row.cells[i].text = str(text)


def op_delete_table_row(doc, op):
    table = _get_table(doc, op.get("table_index", 0))
    row_index = op["row_index"]
    if row_index >= len(table.rows):
        raise ValueError(f"row index {row_index} out of range")
    tr = table.rows[row_index]._tr
    tr.getparent().remove(tr)


def _insert_in_schema_order(parent, new_elem, order):
    """
    OOXML complex types (like CT_TcPr, CT_PPr) require child elements in a
    fixed sequence -- appending blindly at the end can silently produce a
    file real Microsoft Word rejects as "unreadable content" even though
    lenient parsers like python-docx/LibreOffice open it fine. `order` is
    the list of local tag names in required sequence; this inserts
    new_elem immediately before the first existing child that must come
    after it, or appends if none do.
    """
    new_tag = new_elem.tag.split("}")[-1]
    new_idx = order.index(new_tag)
    for child in parent:
        child_tag = child.tag.split("}")[-1]
        if child_tag in order and order.index(child_tag) > new_idx:
            child.addprevious(new_elem)
            return
    parent.append(new_elem)


def op_add_table_column(doc, op):
    """
    python-docx has no native add_column-with-data helper that also extends
    existing rows, so this adds a cell to every existing row manually and
    widens the table's grid definition to match.
    """
    table = _get_table(doc, op.get("table_index", 0))
    width_dxa = int(op.get("width", 3000))  # width is in DXA (twips), same units as columnWidths elsewhere
    cells = op.get("cells", [])
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    new_gridcol = OxmlElement("w:gridCol")
    new_gridcol.set(qn("w:w"), str(width_dxa))  # required attribute -- omitting it produced invalid files
    grid.append(new_gridcol)
    for i, row in enumerate(table.rows):
        new_tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(width_dxa))
        tc_w.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w)
        new_tc.append(tc_pr)
        p = OxmlElement("w:p")
        if i < len(cells):
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = str(cells[i])
            r.append(t)
            p.append(r)
        new_tc.append(p)
        row._tr.append(new_tc)


def op_delete_table_column(doc, op):
    table = _get_table(doc, op.get("table_index", 0))
    col_index = op["col_index"]
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    grid_cols = grid.findall(qn("w:gridCol"))
    if col_index < len(grid_cols):
        grid.remove(grid_cols[col_index])
    for row in table.rows:
        tcs = row._tr.findall(qn("w:tc"))
        if col_index < len(tcs):
            row._tr.remove(tcs[col_index])


def op_merge_cells(doc, op):
    table = _get_table(doc, op.get("table_index", 0))
    r1, c1 = op["start"]
    r2, c2 = op["end"]
    a = table.cell(r1, c1)
    b = table.cell(r2, c2)
    a.merge(b)


_TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
]


def op_set_cell_shading(doc, op):
    table = _get_table(doc, op.get("table_index", 0))
    row_index, col_index = op["cell"]
    color = op["color"].replace("#", "")
    cell = table.cell(row_index, col_index)
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)  # CT_TcPr allows at most one shd -- e.g. a header cell already has one
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    # Must not just append: CT_TcPr requires a fixed child sequence, and this
    # cell may already have e.g. vAlign (which must come AFTER shd) set from
    # creation -- blind append put shd in the wrong place in exactly that case.
    _insert_in_schema_order(tc_pr, shd, _TCPR_ORDER)


def op_set_table_borders(doc, op):
    table = _get_table(doc, op.get("table_index", 0))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # base template lacks this named style; borders just stay default


# ---- Bookmarks ----

_bookmark_id_counter = [1000]


def op_add_bookmark(doc, op):
    anchor = op["anchor"]
    name = op["name"]
    paragraph = _find_paragraph_containing(doc, anchor)
    if paragraph is None:
        raise ValueError(f"could not find paragraph containing: {anchor!r}")
    bm_id = str(_bookmark_id_counter[0])
    _bookmark_id_counter[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bm_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    p_elem = paragraph._p
    # CT_P requires pPr (if present) to be the very first child -- inserting
    # bookmarkStart at index 0 unconditionally put it BEFORE pPr whenever the
    # paragraph had one (e.g. a heading style, or after format_paragraph ran),
    # which is exactly the malformed-order bug real Word rejected as
    # "unreadable content" while python-docx/LibreOffice tolerated it.
    ppr = p_elem.find(qn("w:pPr"))
    insert_index = 1 if ppr is not None else 0
    p_elem.insert(insert_index, start)
    p_elem.append(end)


# ---- Field codes (PAGE, NUMPAGES, TOC) ----

def _add_field(paragraph, field_code):
    """
    Inserts a Word field (e.g. 'PAGE', 'NUMPAGES', 'TOC \\o "1-3" \\h \\z \\u')
    using the standard begin/instrText/separate/end run sequence. Word computes
    the displayed value when the field is updated (F9 / on open with update-fields
    settings) -- this writes the field definition, not a pre-baked static number.
    """
    def make_run(inner_builder):
        r = OxmlElement("w:r")
        inner_builder(r)
        return r

    r1 = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)

    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    r2.append(instr)

    r3 = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fld_sep)

    r4 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r4.append(fld_end)

    for r in (r1, r2, r3, r4):
        paragraph._p.append(r)


def op_insert_page_number_field(doc, op):
    """Adds 'Page X of Y' to the header or footer (default: footer)."""
    target = op.get("target", "footer")
    section = doc.sections[0]
    container = section.footer if target == "footer" else section.header
    p = container.paragraphs[0] if container.paragraphs[0].text == "" else container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    _add_field(p, "PAGE")
    p.add_run(" of ")
    _add_field(p, "NUMPAGES")


def op_insert_toc_field(doc, op):
    """
    Inserts a Table of Contents field. Word computes and displays the actual
    entries when the document is opened (or F9 is pressed) -- python-docx
    cannot pre-render the resolved TOC text itself, only the field definition.
    """
    p = doc.add_paragraph()
    _add_field(p, 'TOC \\o "1-3" \\h \\z \\u')


# ---- Section layout ----

def op_set_orientation(doc, op):
    section = doc.sections[0]
    orientation = op.get("orientation", "portrait")
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def op_set_margins(doc, op):
    section = doc.sections[0]
    if "top" in op:
        section.top_margin = Inches(op["top"])
    if "bottom" in op:
        section.bottom_margin = Inches(op["bottom"])
    if "left" in op:
        section.left_margin = Inches(op["left"])
    if "right" in op:
        section.right_margin = Inches(op["right"])


# ---- Watermark ----

def op_add_watermark(doc, op):
    """
    Adds a diagonal, light-gray text watermark to the header of every section
    -- the standard technique Word itself uses (a WordArt-style textbox placed
    behind text in the header), built here from raw VML since python-docx has
    no built-in watermark API and the VML 'v:' namespace isn't registered in
    python-docx's own OxmlElement namespace map (only OOXML namespaces are),
    so this parses a standalone XML fragment with its own namespace
    declarations instead of using OxmlElement for the VML parts.
    """
    from lxml import etree

    text = op.get("text", "DRAFT")
    color = op.get("color", "C0C0C0").replace("#", "")

    pict_xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:o="urn:schemas-microsoft-com:office:office">
  <v:shapetype id="_x0000_t136" coordsize="1600,21600" o:spt="136"
      adj="10800" path="m@7,0l@8,0m@5,21600l@6,21600e"></v:shapetype>
  <v:shape id="WatermarkShape" type="#_x0000_t136"
      style="position:absolute;left:0;top:0;width:400pt;height:100pt;rotation:315;z-index:-1"
      fillcolor="#{color}" stroked="f">
    <v:fill opacity=".5"/>
    <v:textpath style="font-family:Calibri;font-size:54pt" string="{text}"/>
  </v:shape>
</w:pict>'''
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs[0].text == "" else header.add_paragraph()
        run = p.add_run()
        run._r.append(etree.fromstring(pict_xml))


OP_HANDLERS = {
    "find_replace": op_find_replace,
    "append_paragraph": op_append_paragraph,
    "append_heading": op_append_heading,
    "append_table": op_append_table,
    "append_image": op_append_image,
    "insert_page_break": op_insert_page_break,
    "add_comment": op_add_comment,
    "insert_hyperlink": op_insert_hyperlink,
    "set_header": op_set_header,
    "set_footer": op_set_footer,
    "resize_image": op_resize_image,
    "remove_page": op_remove_page,
    "move_page": op_move_page,
    "set_metadata": op_set_metadata,
    "format_paragraph": op_format_paragraph,
    "add_table_row": op_add_table_row,
    "delete_table_row": op_delete_table_row,
    "add_table_column": op_add_table_column,
    "delete_table_column": op_delete_table_column,
    "merge_cells": op_merge_cells,
    "set_cell_shading": op_set_cell_shading,
    "set_table_borders": op_set_table_borders,
    "add_bookmark": op_add_bookmark,
    "insert_page_number_field": op_insert_page_number_field,
    "insert_toc_field": op_insert_toc_field,
    "set_orientation": op_set_orientation,
    "set_margins": op_set_margins,
    "add_watermark": op_add_watermark,
}


def load_ops(args):
    if args.ops:
        with open(args.ops, "r", encoding="utf-8") as f:
            return json.load(f)
    if args.find is not None and args.replace is not None:
        return [{"op": "find_replace", "find": args.find, "replace": args.replace}]
    return []


def main():
    parser = argparse.ArgumentParser(description="Edit an existing .docx file")
    parser.add_argument("input", help="Path to the source .docx file")
    parser.add_argument("--out", required=True, help="Path to write the edited .docx file")
    parser.add_argument("--ops", default=None, help="Path to a JSON file listing operations")
    parser.add_argument("--find", default=None, help="Shorthand: single find/replace search text")
    parser.add_argument("--replace", default=None, help="Shorthand: single find/replace replacement text")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"error: file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    ops = load_ops(args)
    if not ops:
        print("error: no operations given (use --ops or --find/--replace)", file=sys.stderr)
        sys.exit(1)

    doc = Document(args.input)

    for op in ops:
        op_type = op.get("op")
        handler = OP_HANDLERS.get(op_type)
        if not handler:
            print(f"error: unknown op '{op_type}'", file=sys.stderr)
            sys.exit(1)
        handler(doc, op)

    doc.save(args.out)

    # Self-check: reopen the file we just wrote to confirm it's structurally valid
    try:
        Document(args.out)
    except Exception as e:
        print(f"error: output file failed validation after write: {e}", file=sys.stderr)
        sys.exit(1)

    size = os.path.getsize(args.out)
    print(f"Wrote {args.out} ({size} bytes), applied {len(ops)} operation(s)")


if __name__ == "__main__":
    main()
